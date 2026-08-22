"""Diagnose cross-paper contamination in INSI submission package.
Scan all docx files for AKI-paradox keywords vs spine-paper keywords.
"""
import os, re
from docx import Document

PKG = r"C:\Users\admin\WorkBuddy\2026-07-06-12-34-04\submission\submission_package"

# AKI paradox paper keywords (should NOT appear in INSI spine submission)
AKI_KEYWORDS = [
    "AKI", "acute kidney injury", "CKD", "chronic kidney disease",
    "MIMIC-IV", "MIMIC", "KDIGO", "creatinine", "SOFA", "renal",
    "kidney", "hemodialysis", "ICU mortality", "Stage 3 AKI",
    "eGFR", "nephrology", "dialysis", "diuresis", "oliguria",
]
# Spine paper keywords (expected)
SPINE_KEYWORDS = [
    "spinal cord", "qMRI", "harmonisation", "harmonization", "ComBat",
    "vendor", "spine-generic", "CSA", "MTR", "MTsat", "PERMANOVA",
    "diffusivity", "myelopathy",
]

def extract_text(path):
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs]
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                parts.append(cell.text)
    # headers/footers
    for sec in doc.sections:
        for hf in (sec.header, sec.footer):
            for p in hf.paragraphs:
                parts.append(p.text)
    return "\n".join(parts)

files = []
for root, dirs, names in os.walk(PKG):
    for n in names:
        if n.endswith(".docx") and not n.startswith("~"):
            files.append(os.path.join(root, n))

print(f"Scanning {len(files)} docx files in {PKG}\n")
for f in sorted(files):
    rel = os.path.relpath(f, PKG)
    try:
        text = extract_text(f)
    except Exception as e:
        print(f"[ERROR] {rel}: {e}")
        continue
    lower = text.lower()
    aki_hits = {}
    for kw in AKI_KEYWORDS:
        # word-boundary match, case-insensitive
        n = len(re.findall(r"\b" + re.escape(kw) + r"\b", lower))
        if n > 0:
            aki_hits[kw] = n
    spine_hits = sum(1 for kw in SPINE_KEYWORDS if kw.lower() in lower)
    # First non-empty paragraph as title hint
    first = next((p.strip() for p in text.split("\n") if p.strip()), "")[:100]
    print(f"=== {rel} ({len(text)} chars) ===")
    print(f"  First line: {first}")
    print(f"  Spine keywords present: {spine_hits}/{len(SPINE_KEYWORDS)}")
    if aki_hits:
        print(f"  !! AKI-related keyword hits: {aki_hits}")
    else:
        print(f"  AKI keywords: none")
    print()
